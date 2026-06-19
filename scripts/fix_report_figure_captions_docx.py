#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""修正 Word 报告中图8–图18 标题与正文引用，仅改文字、不改样式布局。"""

from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
except ImportError:
    print("请先安装: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "数据挖掘实训报告.docx"

# 精确全文替换（先执行）
TEXT_REPLACEMENTS = [
    ("图8 训练/验证Loss与Acc曲线", "图9 训练与验证 Loss、Acc 曲线"),
    ("图12 错例", "图12 测试集高置信度错分样例"),
    ("图14：根目录ONNX与Android assets/models路径对照", "图14 ONNX 模型路径对照"),
    ("图15 工程结构与模型资产", "图15 工程结构与模型资源路径"),
    ("图6 训练集各类别样本频数柱状图", "图6 训练集各类别样本频数分布"),
    (
        "图18将第5章 training_config.json 中的 best_val_acc=61.55、num_classes=50 展示给用户；"
        "输入尺寸 224 与第7.1节预处理一致；推理线程数由 ClassifierProvider.getThreadCount() 读取，可按设备配置。",
        "图18展示验证准确率61.55%、分类数50、输入尺寸224等训练配置信息；"
        "输入尺寸与7.1节预处理一致，推理线程数可按设备配置。",
    ),
]


def set_paragraph_text(paragraph, text: str) -> None:
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def copy_paragraph_style(source, target) -> None:
    target.style = source.style
    target.alignment = source.alignment
    if source.runs and target.runs:
        src_run = source.runs[0]
        dst_run = target.runs[0]
        dst_run.bold = src_run.bold
        dst_run.italic = src_run.italic
        dst_run.font.name = src_run.font.name
        dst_run.font.size = src_run.font.size


def insert_paragraph_after(reference, text: str, style_ref=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    reference._p.addnext(new_p)
    new_para = Paragraph(new_p, reference._parent)
    if style_ref is not None:
        copy_paragraph_style(style_ref, new_para)
    new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def dedupe_figure_blocks(doc: Document) -> int:
    """删除连续重复的图标题+说明段落对。"""
    removed = 0
    i = 0
    while i < len(doc.paragraphs) - 1:
        a = doc.paragraphs[i].text.strip()
        b = doc.paragraphs[i + 1].text.strip()
        if (
            doc.paragraphs[i].style.name == "图标题"
            and a.startswith("图")
            and doc.paragraphs[i + 1].style.name != "图标题"
            and b.startswith(a.split()[0])
        ):
            j = i + 2
            if j + 1 < len(doc.paragraphs):
                c = doc.paragraphs[j].text.strip()
                d = doc.paragraphs[j + 1].text.strip()
                if doc.paragraphs[j].style.name == "图标题" and c == a and d == b:
                    remove_paragraph(doc.paragraphs[j + 1])
                    remove_paragraph(doc.paragraphs[j])
                    removed += 2
                    continue
        i += 1
    return removed


def figure_block_exists(doc: Document, caption: str) -> bool:
    return any(p.text.strip() == caption and p.style.name == "图标题" for p in doc.paragraphs)


def find_paragraph(doc: Document, substring: str, start: int = 0) -> tuple[int, Paragraph] | None:
    for i in range(start, len(doc.paragraphs)):
        if substring in doc.paragraphs[i].text:
            return i, doc.paragraphs[i]
    return None


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    set_paragraph_text(paragraph, paragraph.text.replace(old, new))
    return True


def main() -> None:
    if not DOCX_PATH.exists():
        print(f"找不到: {DOCX_PATH}")
        sys.exit(1)

    doc = Document(str(DOCX_PATH))
    changed = 0

    n_dup = dedupe_figure_blocks(doc)
    if n_dup:
        changed += n_dup
        print(f"去重: 删除 {n_dup} 个重复段落")

    for old, new in TEXT_REPLACEMENTS:
        for p in doc.paragraphs:
            if replace_in_paragraph(p, old, new):
                changed += 1
                print(f"替换: {old[:36]}...")

    # 图标题 / 正文样式参照
    cap_ref = find_paragraph(doc, "图7 ResNet18")[1]
    body_ref = find_paragraph(doc, "图7中，反向传播")[1]

    # --- 5.4 节：补 图8、图10 ---
    _, train_summary = find_paragraph(doc, "训练32 epoch后早停")
    _, wrong_cap = find_paragraph(doc, "图9 训练与验证 Loss、Acc 曲线")
    _, fig9_body = find_paragraph(doc, "图9展示表8中")

    # 空段落用作 图8 标题
    empty_after_summary = doc.paragraphs[find_paragraph(doc, "训练32 epoch后早停")[0] + 1]
    if not figure_block_exists(doc, "图8 训练过程 Epoch/Loss/Acc 日志"):
        if not empty_after_summary.text.strip():
            copy_paragraph_style(cap_ref, empty_after_summary)
            set_paragraph_text(empty_after_summary, "图8 训练过程 Epoch/Loss/Acc 日志")
            insert_paragraph_after(
                empty_after_summary,
                "图8显示 Epoch 22 附近验证 Acc 达到峰值，之后出现震荡或回落，支撑早停决策。",
                body_ref,
            )
            changed += 2
            print("补全: 图8 标题与说明")
        else:
            print("跳过 图8：无空段落可写入")
    else:
        print("已有 图8 标题")

    if not figure_block_exists(doc, "图10 训练配置文件与模型权重文件信息"):
        insert_paragraph_after(
            fig9_body,
            "图10 训练配置文件与模型权重文件信息",
            cap_ref,
        )
        insert_paragraph_after(
            find_paragraph(doc, "图10 训练配置文件与模型权重文件信息")[1],
            "图10保证实验可追溯：配置JSON与权重文件修改时间对应同一次训练 run。",
            body_ref,
        )
        changed += 2
        print("补全: 图10 标题与说明")
    else:
        print("已有 图10 标题")

    # --- 6.5 节：补 图13 ---
    if not figure_block_exists(doc, "图13 ONNX 与 PyTorch 推理一致性验证输出"):
        _, logits_p = find_paragraph(doc, "logits 最大绝对差")
        cap13 = insert_paragraph_after(
            logits_p,
            "图13 ONNX 与 PyTorch 推理一致性验证输出",
            cap_ref,
        )
        insert_paragraph_after(
            cap13,
            "图13表明，部署模型与训练权重在真实样本上推理结果一致，优于仅使用随机张量进行的冒烟测试。",
            body_ref,
        )
        changed += 2
        print("补全: 图13 标题与说明")
    else:
        print("已有 图13 标题")

    # --- 图15 说明 ---
    _, fig15_cap = find_paragraph(doc, "图15 工程结构与模型资源路径")
    next_p = doc.paragraphs[find_paragraph(doc, "图15 工程结构与模型资源路径")[0] + 1]
    if not next_p.text.strip():
        set_paragraph_text(next_p, "图15列出与模型部署直接相关的路径：")
        copy_paragraph_style(body_ref, next_p)
        changed += 1
        print("补全: 图15 说明")

    # 图12 标题样式对齐
    _, fig12_cap = find_paragraph(doc, "图12 测试集高置信度错分样例")
    if fig12_cap.style.name != cap_ref.style.name:
        copy_paragraph_style(cap_ref, fig12_cap)
        print("修正: 图12 标题样式")

    # 图6 标题样式与 md 文案
    hit = find_paragraph(doc, "图6 训练集各类别样本频数")
    if hit:
        _, fig6_cap = hit
        if fig6_cap.style.name != cap_ref.style.name:
            copy_paragraph_style(cap_ref, fig6_cap)
            print("修正: 图6 标题样式")

    doc.save(str(DOCX_PATH))
    print(f"\n完成，共 {changed} 处修改。已保存: {DOCX_PATH}")

    # 验证图1–图18 标题
    print("\n=== 图标题核查 ===")
    import re

    found: dict[int, str] = {}
    for p in doc.paragraphs:
        t = p.text.strip()
        m = re.match(r"^图(\d+)\s+(.+)$", t)
        if m and p.style.name == "图标题" and len(m.group(2)) < 60:
            n = int(m.group(1))
            if n not in found:
                found[n] = t

    for n in range(1, 19):
        if n in found:
            print(f"  OK  {found[n]}")
        else:
            print(f"  缺失 图{n}")


if __name__ == "__main__":
    main()
