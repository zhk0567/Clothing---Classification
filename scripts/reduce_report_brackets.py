#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""降低报告括号与反引号密度。"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "数据挖掘实训报告.md"
DOCX_PATH = ROOT / "数据挖掘实训报告.docx"

REPLACEMENTS = [
    # 摘要
    (
        "前期对本地`服装/`目录（11287张、3个粗粒度子类）进行量化探索，并与DeepFashion官方基准（20000张、50细类、规范标注）对比，",
        "前期对本地服装目录共11287张、3个粗粒度子类进行量化探索，并与DeepFashion官方基准20000张、50个细类、规范标注对比，",
    ),
    (
        "建模阶段采用ImageNet预训练ResNet18进行**全网络微调**（非冻结骨干），配合标签平滑、AdamW、余弦退火学习率与Early Stopping。优化后模型在验证集取得最佳Top-1准确率61.55%（Epoch 22，独立评估脚本复现为61.50%，差异小于0.1个百分点），Top-3为79.75%；独立测试集Top-1为58.73%（Top-3为79.00%）。Android端推理已对齐训练预处理（Resize 256 + CenterCrop 224），消除Training-Serving Skew。",
        "建模阶段采用ImageNet预训练ResNet18进行**全网络微调**，卷积骨干未冻结，并配合标签平滑、AdamW、余弦退火学习率与Early Stopping。"
        "优化后模型在验证集Epoch 22取得最佳Top-1准确率61.55%，独立评估脚本复现为61.50%；Top-3准确率为79.75%。"
        "独立测试集Top-1为58.73%，Top-3为79.00%。Android端推理已与训练预处理对齐，采用Resize 256与CenterCrop 224，消除Training-Serving Skew。",
    ),
    (
        "本报告重点给出各环节方法原理、核心代码与实测评估结果（含混淆矩阵与错例分析）；",
        "本报告重点给出各环节方法原理、核心代码与实测评估结果，含混淆矩阵与错例分析；",
    ),
    # 1.2
    (
        "y_i 为50类标签之一（0～49），",
        "y_i 为50类标签之一，编号0～49，",
    ),
    (
        "项目含完整脚本链（`train_deepfashion_complete.py`、`evaluate_deepfashion_model.py`、`update_model_for_android.py`）；",
        "项目含训练、评估与模型导出等完整脚本；",
    ),
    # 表题与节标题
    ("**表2 本地`服装/`目录统计（`python scripts/stats_local_clothing.py`）**", "**表2 本地服装目录统计**"),
    ("### 3.4 类别不平衡（实测统计）", "### 3.4 类别不平衡实测"),
    ("### 4.1 传统方法与深度方法对比（概念+课设选型）", "### 4.1 传统方法与深度方法对比"),
    ("**表5 传统方法与深度学习方法对比（课设选型）**", "**表5 传统方法与深度学习方法对比**"),
    ("### 4.4 训练/验证循环（完整准确率计算）", "### 4.4 训练/验证循环"),
    ("**表6 训练环境（`train_deepfashion_complete.py`）**", "**表6 训练环境**"),
    ("**表7 评估复现环境（`evaluate_deepfashion_model.py`）**", "**表7 评估复现环境**"),
    ("**表10 验证集与测试集评估结果（`evaluation_results.json`）**", "**表10 验证集与测试集评估结果**"),
    ("**表11 宏平均指标说明（测试集）**", "**表11 宏平均指标说明**"),
    ("**表12 测试集Top10易混类别对（真实→预测）**", "**表12 测试集Top10易混类别对**"),
    ("**表13 测试集F1较高类别（support≥50）**", "**表13 测试集F1较高类别**"),
    ("**表14 测试集F1较低类别（support≥10）**", "**表14 测试集F1较低类别**"),
    ("### 6.4 高置信度错例（softmax概率）", "### 6.4 高置信度错例分析"),
    ("**表19 复现命令（项目根目录 PowerShell）**", "**表19 复现命令**"),
    # 表格单元格
    ("| 合计 | 20000 | 20000 | 50（全集） |", "| 合计 | 20000 | 20000 | 50，全部类别 |"),
    ("| 类别数 | 3（粗类） | 50（细类） |", "| 类别数 | 3类粗粒度 | 50类细粒度 |"),
    ("| 未单独实现（细粒度效果预期差） |", "| 未单独实现，细粒度效果预期较差 |"),
    ("| 未采用（收敛慢、易过拟合） |", "| 未采用，收敛慢且易过拟合 |"),
    ("| 初版（过拟合） | 60.90% |", "| 初版方案 | 60.90% |"),
    ("| 优化后（部署版） | **61.55%**（Ep.22） |", "| 优化后方案 | **61.55%**，Epoch 22 |"),
    ("| max_epochs | 100（实际32） | 上限保护 |", "| max_epochs | 100，实际32轮 | 上限保护 |"),
    ("| PyTorch | CUDA版（训练时） |", "| PyTorch | CUDA版，用于训练 |"),
    ("| 设备 | CPU（结果与GPU推理一致） |", "| 设备 | CPU，与GPU推理结果一致 |"),
    # 正文括号
    ("表3说明：全集50类，但单一划分中未必包含全部类别（如训练集45类），", "表3说明：全集50类，但单一划分中未必包含全部类别，例如训练集仅含45类，"),
    (
        "**图2** DeepFashion 数据集目录结构（`Anno_fine` 与 `Img/img`）",
        "**图2** DeepFashion 数据集目录结构",
    ),
    ("**图4** DeepFashion 原始样本示例（Tee、Jeans、Dress）", "**图4** DeepFashion 原始样本示例"),
    ("感知哈希（pHash）去重", "感知哈希pHash去重"),
    ("最多4016张/类（Dress），", "最多4016张/类，Dress类，"),
    ("最多4016张（Dress），", "最多4016张，Dress类，"),
    ("Dress占1064/4000（26.6%），", "Dress占1064/4000，约26.6%，"),
    (
        "本实训**未冻结**卷积骨干，而是加载ImageNet预训练权重后，**替换分类头并对全部参数进行微调**（`model.parameters()`全部进入AdamW）。",
        "本实训**未冻结**卷积骨干，而是加载ImageNet预训练权重后，**替换分类头并对全部参数进行微调**，全部网络参数纳入AdamW优化。",
    ),
    ("Early Stopping是模型选择（Model Selection）标准技术：", "Early Stopping是模型选择标准技术："),
    (
        "训练32 epoch后早停；Epoch 22验证集最佳61.55%（`training_config.json`）；",
        "训练32 epoch后早停；Epoch 22验证集最佳Top-1为61.55%；",
    ),
    (
        "本章使用`scripts/evaluate_deepfashion_model.py`在**独立测试集4000张**上评估，",
        "本章在独立测试集4000张样本上评估，",
    ),
    (
        "表10说明：验证集指标以训练时`training_config.json`记录的最佳值61.55%为准；评估脚本复现为61.50%。测试集Top-1低2.77个百分点，属正常泛化落差。Top-3均为79%左右，应用 Top-3 展示与测试集 Top-3 准确率 79.00% 的评估结果一致。随机Top-1基线2%（1/50）。",
        "表10说明：验证集最佳Top-1以训练记录61.55%为准，评估脚本复现为61.50%。测试集Top-1低2.77个百分点，属正常泛化落差。"
        "Top-3准确率约为79%，与应用中Top-3展示一致。随机猜测Top-1基线约为2%。",
    ),
    (
        "**报告以Top-1与“有样本类”宏F1（0.221）为主**。",
        "**报告以Top-1与有样本类的宏F1 0.221为主**。",
    ),
    (
        "模型对**连衣裙（Dress）存在预测偏好**，多条样本被误判为 Dress。",
        "模型对**Dress连衣裙**存在预测偏好，多条样本被误判为Dress。",
    ),
    (
        "**Dress类样本量最大（训练4016/测试1064），模型对其有系统性过预测倾向**，是后续校准（阈值/重采样）的重点。",
        "**Dress类样本量最大，训练集4016张、测试集1064张，模型对其存在系统性过预测倾向**，后续可通过阈值调整或重采样校准。",
    ),
    (
        "文件名含真实/预测类与**softmax置信度**（非logit）。",
        "文件名含真实类、预测类及softmax置信度。",
    ),
    (
        "典型错例：Romper↔Jumpsuit（连体装类别相互混淆）、Tee→Tank、Blouse→Jumpsuit 等。",
        "典型错例包括Romper与Jumpsuit、Tee与Tank、Blouse与Jumpsuit等类别相互混淆。",
    ),
    (
        "`scripts/test_onnx_model.py`使用**真实测试图像**（非随机噪声），对比PyTorch与ONNX的Top-1及logits差异：",
        "使用真实测试图像对比PyTorch与ONNX的Top-1结果及logits差异：",
    ),
    (
        "Android 端采用与训练端一致的 `Resize(256)+CenterCrop(224)` 预处理（`DeepFashionClassifier.kt`），避免 Training-Serving Skew：",
        "Android端采用与训练端一致的Resize 256与CenterCrop 224预处理，在DeepFashionClassifier中实现，以避免Training-Serving Skew：",
    ),
    (
        "归一化仍使用 ImageNet 均值方差，与 Python `Normalize` 一致。",
        "归一化仍使用ImageNet均值方差，与Python端Normalize一致。",
    ),
    # 第7章精简
    (
        "**图15 工程结构与模型资产**\n\n**图15** Android 工程结构与模型资源路径",
        "**图15** 工程结构与模型资源路径",
    ),
    (
        "**图16 识别结果页 ResultActivity**\n\n**图16** 识别结果页（ResultActivity）",
        "**图16** 识别结果页",
    ),
    (
        "**图17 数据统计页 StatisticsActivity**\n\n**图17** 数据统计页（StatisticsActivity）",
        "**图17** 数据统计页",
    ),
    (
        "**图18 模型信息页 ModelInfoActivity**\n\n**图18** 模型信息页（ModelInfoActivity）",
        "**图18** 模型信息页",
    ),
    (
        "**入口**：相机单拍（`CameraActivity`）、相册（`ClassifyFragment`）、裁剪重识别（`ImageCropActivity`）、增强重识别（`ImageEnhanceActivity`）、历史记录回看。",
        "**入口**：相机拍照、相册选图、裁剪后重识别、增强后重识别及历史记录回看。",
    ),
    ("**页面要素（与代码一致）**：", "**页面要素**："),
    (
        "1. **Top-1 主结果**：`tvCategory` 显示中文类名（如「T恤」），`tvConfidence` 与进度条显示 Softmax 置信度百分比。",
        "1. **Top-1 主结果**：显示中文类名及Softmax置信度百分比与进度条。",
    ),
    (
        "2. **Top-3 切换**：`ClassifierProvider.classifyImageTopK(..., k=3)` 返回前三候选；`TabLayout` 标签为 `#1 85%`、`#2 12%` 等，切换 Tab 可查看不同候选的类别与百科，对应测试集 **Top-3 准确率 79.00%** 的界面展示。",
        "2. **Top-3 切换**：展示前三候选类别、置信度及百科信息，与测试集Top-3准确率79.00%相对应。",
    ),
    (
        "4. **类别百科卡片**：风格、场景、搭配、季节/材质、护理等字段，数据来自 `CategoryRepository` 预置描述，将「类别 ID」转化为可读知识。",
        "4. **类别百科卡片**：展示风格、场景、搭配、季节材质与护理等预置描述信息。",
    ),
    (
        "6. **操作菜单**：分享（文字/带水印图片）、裁剪重识别、增强重识别；新结果可写入 `history.jsonl`。",
        "6. **操作菜单**：支持分享、裁剪重识别、增强重识别，并将新结果写入本地历史记录。",
    ),
    (
        "**入口**：底部导航「更多」→ `StatisticsActivity`（`MoreFragment.cardStatistics`）。",
        "**入口**：底部导航「更多」进入数据统计页。",
    ),
    ("**表16 StatisticsActivity 页面要素与数据来源**", "**表16 数据统计页要素与数据来源**"),
    (
        "上述统计均从本地 `history.jsonl` 读取，**不上传云端**，属于基于本地使用记录的统计分析，可补充考察识别类别分布、使用频次与平均置信度等指标，与第6章基于固定测试集的离线评估形成补充。",
        "上述统计均从本地history.jsonl读取且不上传云端，属于基于本地使用记录的统计分析，可补充考察识别类别分布、使用频次与平均置信度，与第6章离线测试集评估相互补充。",
    ),
    (
        "**入口**：「更多」→ `ModelInfoActivity`（`MoreFragment.cardModelInfo`）。",
        "**入口**：「更多」进入模型信息页。",
    ),
    (
        "**页面固定展示字段**（`strings.xml` → `model_info_content`，与训练配置对照）：",
        "**页面固定展示字段**如下，与训练配置一致：",
    ),
    (
        "推理线程：N（随 CPU 核心数/设置变化）",
        "推理线程：随CPU核心数或设置变化",
    ),
    (
        "图18将第5章 `training_config.json` 中的 `best_val_acc=61.55`、`num_classes=50` 展示给用户；输入尺寸 224 与第7.1节预处理一致；推理线程数由 `ClassifierProvider.getThreadCount()` 读取，可按设备配置。",
        "图18展示验证准确率61.55%、分类数50、输入尺寸224等训练配置信息；输入尺寸与7.1节预处理一致，推理线程数可按设备配置。",
    ),
    (
        "DeepFashionClassifier → ResultActivity（Top-1/Top-3）",
        "DeepFashionClassifier → 识别结果页，展示Top-1与Top-3",
    ),
    (
        "history.jsonl → StatisticsActivity（趋势/分布）",
        "history.jsonl → 数据统计页，展示趋势与分布",
    ),
    (
        "ModelInfoActivity（模型版本与准确率）",
        "模型信息页，展示模型版本与准确率",
    ),
    (
        "识别结果写入 `history.jsonl` 后，可在 StatisticsActivity 查看使用趋势与类别分布，在 ModelInfoActivity 查看模型版本与验证准确率。",
        "识别结果写入本地历史记录后，可在数据统计页查看使用趋势与类别分布，在模型信息页查看模型版本与验证准确率。",
    ),
    # 第8章
    (
        "ResNet18 全网络微调取得验证 61.55%（复现 61.50%）、测试 58.73%；",
        "ResNet18全网络微调取得验证Top-1 61.55%、复现61.50%、测试58.73%；",
    ),
    (
        "（1）**数据决策**：本地11287张粗类数据不可直接用于50类任务；DeepFashion为唯一训练源。（2）**模型规律**：Dress样本最多，模型对其有过预测倾向；Skirt/Top/Blouse等细类是主要混淆源。（3）**部署要点**：预处理必须与训练一致；ONNX在真实样本上与PyTorch logits差小于1e-5。",
        "**数据决策**：本地11287张粗类数据不可直接用于50类任务，DeepFashion为唯一训练源。\n\n**模型规律**：Dress样本最多，模型对其有过预测倾向；Skirt、Top、Blouse等细类是主要混淆源。\n\n**部署要点**：预处理必须与训练一致；ONNX在真实样本上与PyTorch logits差小于1e-5。",
    ),
    ("*封面【】请填写。*", "*封面信息请填写。*"),
    # 其他正文反引号
    ("预处理在`scripts/train_deepfashion_complete.py`实现，", "预处理在训练脚本中实现，"),
    ("项目早期建立本地目录`服装/`，", "项目早期建立本地服装目录，"),
    ("**图1** 本地`服装/`目录结构", "**图1** 本地服装目录结构"),
    ("正式数据源为`Category and Attribute Prediction Benchmark/`。", "正式数据源为Category and Attribute Prediction Benchmark。"),
    ("应排查压缩包是否完整解压至`Img/img/`。", "应排查压缩包是否完整解压至Img/img目录。"),
    ("| 维度 | 本地`服装/` | DeepFashion |", "| 维度 | 本地服装目录 | DeepFashion |"),
    ("| 数据理解 | 本地`服装/`统计、DeepFashion元数据核对 |", "| 数据理解 | 本地服装目录统计、DeepFashion元数据核对 |"),
    ("本地`服装/`补充50类标注后做域适应。", "本地服装目录补充50类标注后做域适应。"),
    ("对`train_cate.txt`统计：", "对训练集类别标注文件统计："),
    ("表3说明：全集50类，但单一划分中未必包含全部类别，例如训练集仅含45类，属官方划分特性；评估时需按各类support分别解读。", "表3说明：全集50类，但单一划分中未必包含全部类别，例如训练集仅含45类，属官方划分特性；评估时需按各类别support分别解读。"),
    ("运行`python scripts/export_misclassified_samples.py`可导出8张错例至`docs/misclassified_samples/`，", "运行错例导出脚本可导出8张错例，"),
]

DOCX_EXTRA = [
    ("表2 本地服装/目录统计（python scripts/stats_local_clothing.py）", "表2 本地服装目录统计"),
    (
        "训练32 epoch后早停；Epoch 22验证集最佳61.55%（training_config.json）；",
        "训练32 epoch后早停；Epoch 22验证集最佳Top-1为61.55%；",
    ),
    ("表11 宏平均指标说明（测试集）", "表11 宏平均指标说明"),
    ("报告以Top-1与“有样本类”宏F1（0.221）为主。", "报告以Top-1与有样本类的宏F1 0.221为主。"),
    ("表12 测试集Top10易混类别对（真实→预测）", "表12 测试集Top10易混类别对"),
    (
        "Dress类样本量最大（训练4016/测试1064），模型对其有系统性过预测倾向，是后续校准（阈值/重采样）的重点。",
        "Dress类样本量最大，训练集4016张、测试集1064张，模型对其存在系统性过预测倾向，后续可通过阈值调整或重采样校准。",
    ),
    ("表13 测试集F1较高类别（support≥50）", "表13 测试集F1较高类别"),
    ("表14 测试集F1较低类别（support≥10）", "表14 测试集F1较低类别"),
    ("softmax置信度（非logit）", "softmax置信度"),
    ("使用真实测试图像（非随机噪声），", "使用真实测试图像，"),
    (
        "1. Top-1 主结果：tvCategory 显示中文类名（如「T恤」），tvConfidence 与进度条显示 Softmax 置信度百分比。",
        "1. Top-1 主结果：显示中文类名及Softmax置信度百分比与进度条。",
    ),
    (
        "6. 操作菜单：分享（文字/带水印图片）、裁剪重识别、增强重识别；新结果可写入 history.jsonl。",
        "6. 操作菜单：支持分享、裁剪重识别、增强重识别，并将新结果写入本地历史记录。",
    ),
    (
        "页面固定展示字段（strings.xml → model_info_content，与训练配置对照）：",
        "页面固定展示字段如下，与训练配置一致：",
    ),
    ("取得验证61.55%（复现61.50%）、测试58.73%", "取得验证Top-1 61.55%、复现61.50%、测试58.73%"),
    (
        "表10说明：验证集指标以训练时training_config.json记录的最佳值61.55%为准；评估脚本复现为61.50%。测试集Top-1低2.77个百分点，属正常泛化落差。Top-3均为79%左右，应用 Top-3 展示与测试集 Top-3 准确率 79.00% 的评估结果一致。随机Top-1基线2%（1/50）。",
        "表10说明：验证集最佳Top-1以训练记录61.55%为准，评估脚本复现为61.50%。测试集Top-1低2.77个百分点，属正常泛化落差。Top-3准确率约为79%，与应用中Top-3展示一致。随机猜测Top-1基线约为2%。",
    ),
    (
        "2. Top-3 切换：ClassifierProvider.classifyImageTopK(..., k=3) 返回前三候选；TabLayout 标签为 #1 85%、#2 12% 等，切换 Tab 可查看不同候选的类别与百科，对应测试集 Top-3 准确率 79.00% 的界面展示。",
        "2. Top-3 切换：展示前三候选类别、置信度及百科信息，与测试集Top-3准确率79.00%相对应。",
    ),
    (
        "项目含完整脚本链（train_deepfashion_complete.py、evaluate_deepfashion_model.py、update_model_for_android.py）；",
        "项目含训练、评估与模型导出等完整脚本；",
    ),
    ("预处理在scripts/train_deepfashion_complete.py实现，", "预处理在训练脚本中实现，"),
    ("（model.parameters()全部进入AdamW）", "，全部网络参数纳入AdamW优化"),
    (
        "入口：相机单拍（CameraActivity）、相册（ClassifyFragment）、裁剪重识别（ImageCropActivity）、增强重识别（ImageEnhanceActivity）、历史记录回看。",
        "入口：相机拍照、相册选图、裁剪后重识别、增强后重识别及历史记录回看。",
    ),
]


def strip_prose_backticks(text: str) -> str:
    """去掉正文段落中的反引号，保留代码块与表格内路径。"""

    def repl_codeblock(match: re.Match) -> str:
        return match.group(0)

    parts: list[str] = []
    last = 0
    for m in re.finditer(r"```[\s\S]*?```", text):
        parts.append(_strip_inline_ticks(text[last : m.start()]))
        parts.append(m.group(0))
        last = m.end()
    parts.append(_strip_inline_ticks(text[last:]))
    return "".join(parts)


def _strip_inline_ticks(chunk: str) -> str:
    lines = chunk.splitlines(keepends=True)
    out: list[str] = []
    for line in lines:
        if line.lstrip().startswith("|") or line.lstrip().startswith("#"):
            out.append(line)
            continue
        if "http://" in line or "https://" in line:
            out.append(line)
            continue
        out.append(re.sub(r"`([^`\n]+)`", r"\1", line))
    return "".join(out)


def replace_paragraph(paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    t = paragraph.text.replace(old, new)
    runs = paragraph.runs
    if runs:
        runs[0].text = t
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(t)
    return True


def polish_docx(replacements: list[tuple[str, str]]) -> int:
    if Document is None or not DOCX_PATH.exists():
        return 0
    doc = Document(str(DOCX_PATH))
    total = 0
    all_repl = replacements + DOCX_EXTRA
    for old, new in all_repl:
        if not old:
            continue
        for paragraph in doc.paragraphs:
            if replace_paragraph(paragraph, old, new):
                total += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if replace_paragraph(paragraph, old, new):
                            total += 1
    doc.save(str(DOCX_PATH))
    return total


def main() -> None:
    if not MD_PATH.exists():
        print(f"找不到: {MD_PATH}")
        sys.exit(1)

    text = MD_PATH.read_text(encoding="utf-8")
    count = 0
    for old, new in REPLACEMENTS:
        if old in text:
            text = text.replace(old, new)
            count += 1

    text = strip_prose_backticks(text)
    MD_PATH.write_text(text, encoding="utf-8")
    print(f"Markdown 更新 {count} 组替换 → {MD_PATH}")

    n = polish_docx(REPLACEMENTS)
    print(f"Word 更新 {n} 处 → {DOCX_PATH}")


if __name__ == "__main__":
    main()
