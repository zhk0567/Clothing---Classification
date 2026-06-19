#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""仅替换 Word 报告中的修订痕迹类文字，不改动样式。"""

from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("请先安装: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "数据挖掘实训报告.docx"

REPLACEMENTS = [
    (
        "实际训练管线中已实现的清洗包括：路径存在性校验、损坏图像异常捕获。"
        "README中规划的JPEG统一转换、感知哈希（pHash）去重属于数据治理扩展方案，"
        "可在独立预处理脚本中执行，当前训练脚本未内置pHash逻辑——报告如实记录，避免“写了但未实现”。",
        "训练管线已实现的清洗包括：路径存在性校验、损坏图像异常捕获。"
        "JPEG 统一转换、感知哈希（pHash）去重可作为数据治理扩展方案在独立预处理脚本中执行，"
        "当前训练脚本未内置 pHash 逻辑。",
    ),
    (
        "图11对角线为正确分类。修复评估脚本参数顺序后，易混对反映真实误判方向："
        "模型对连衣裙（Dress）存在预测偏好，多条样本被误判为Dress。",
        "图11对角线为正确分类。易混对反映主要误判方向："
        "模型对连衣裙（Dress）存在预测偏好，多条样本被误判为 Dress。",
    ),
    ("3.2 清洗策略说明（如实表述）", "3.2 清洗策略说明"),
    ("3.1 清洗策略说明（如实表述）", "3.1 清洗策略说明"),
    ("4.2 迁移学习与全网络微调（纠正表述）", "4.2 迁移学习与全网络微调"),
    ("7.1 部署链路与训练—服务一致性（已修复）", "7.1 部署链路与训练—服务一致性"),
    (
        "早期Android版本曾直接将图像拉伸为224×224，与训练端Resize(256)+CenterCrop(224)不一致，"
        "会造成Training-Serving Skew。现已在DeepFashionClassifier.kt中修复：",
        "Android 端采用与训练端一致的 Resize(256)+CenterCrop(224) 预处理（DeepFashionClassifier.kt），"
        "避免 Training-Serving Skew：",
    ),
    (
        "归一化仍使用ImageNet均值方差，与Python Normalize一致。修复后，移动端推理与验证集评估采用同一几何变换。",
        "归一化仍使用 ImageNet 均值方差，与 Python Normalize 一致。"
        "移动端推理与验证集评估因此采用同一几何变换。",
    ),
    ("修复Android预处理与评估脚本后，混淆矩阵揭示", "混淆矩阵揭示"),
    ("表20 报告插图文件（已生成可直插）", "表20 报告插图文件"),
    ("表A-4 报告插图文件（已生成可直插）", "表A-4 报告插图文件"),
    (
        "封面【】请填写。图3-3/6-1/6-2可使用docs/下已生成PNG/JPG；其余截图按占位说明补充。",
        "封面【】请填写。",
    ),
    (
        "封面【】请填写。图6/11/12可使用docs/下已生成PNG/JPG；其余截图按占位说明补充。",
        "封面【】请填写。",
    ),
]


def replace_paragraph(paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    text = paragraph.text.replace(old, new)
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    return True


def iter_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is None:
                continue
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph


def main() -> None:
    if not DOCX_PATH.exists():
        print(f"找不到: {DOCX_PATH}")
        sys.exit(1)

    doc = Document(str(DOCX_PATH))
    changed = 0
    for old, new in REPLACEMENTS:
        for paragraph in iter_paragraphs(doc):
            if replace_paragraph(paragraph, old, new):
                changed += 1
                print(f"已替换: {old[:40]}...")

    doc.save(str(DOCX_PATH))
    print(f"完成，共 {changed} 处。已保存: {DOCX_PATH}")


if __name__ == "__main__":
    main()
