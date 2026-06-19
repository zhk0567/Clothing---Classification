#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""润色报告正文：去除口语化表述（保留第五类轻微用语）。"""

from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "数据挖掘实训报告.md"
DOCX_PATH = ROOT / "数据挖掘实训报告.docx"

# (old, new) — 顺序无关时逐条 replace
REPLACEMENTS = [
    (
        "表1表明，本课设并非仅“训练一个神经网络”，而是完整走通数据挖掘标准方法论；深度学习是建模阶段的算法选择，而非替代数据理解与评估环节。",
        "表1表明，本课设覆盖数据挖掘全流程各环节，而非局限于单一建模步骤；深度学习是建模阶段的算法选择，并不替代数据理解与评估环节。",
    ),
    (
        "RTX 4060可加速训练。综合可行。",
        "RTX 4060可加速训练，实验条件具备。",
    ),
    (
        "11287张规模不小，但相对DeepFashion仍缺少",
        "11287张样本量少于DeepFashion，且相对后者仍缺少",
    ),
    (
        "避免“Garbage In, Garbage Out”。",
        "避免因数据质量不足导致模型效果下降。",
    ),
    (
        "【插入截图 图1：本地服装目录文件结构（资源管理器展开`服装/`，显示上衣/下装/连身装及样例缩略图）】",
        "**图1** 本地`服装/`目录结构",
    ),
    (
        "【插入截图 图2：DeepFashion目录结构（`Anno_fine`与`Img/img`）】",
        "**图2** DeepFashion 数据集目录结构（`Anno_fine` 与 `Img/img`）",
    ),
    (
        "【插入截图 图3：PowerShell统计`train.txt`/`val.txt`/`test.txt`行数输出】",
        "**图3** 训练/验证/测试集标注文件行数统计结果",
    ),
    (
        "【插入截图 图4：DeepFashion原始样本（Tee、Jeans、Dress等不同类别各1张）】",
        "**图4** DeepFashion 原始样本示例（Tee、Jeans、Dress）",
    ),
    (
        "【插入截图 图5：同一张原图经增强后的对比拼图】",
        "**图5** 数据增强前后对比",
    ),
    (
        "图5展示增强如何“免费”扩充样本空间，是控制过拟合的数据侧手段。",
        "图5表明，数据增强在不增加原始标注成本的前提下扩充有效样本空间，是控制过拟合的数据侧手段。",
    ),
    (
        "【插入截图 图6：训练集各类别样本频数柱状图（运行`python scripts/plot_class_distribution.py`生成`docs/train_class_distribution.png`）】",
        "**图6** 训练集各类别样本频数分布",
    ),
    (
        "【插入截图 图7：ResNet18全网络微调结构示意（预训练骨干+Dropout+FC512+FC50）】",
        "**图7** ResNet18 全网络微调结构示意",
    ),
    (
        "| 初版（过拟合） | 60.90% | 明显更高 | **>27%** | Dropout偏低、增强较弱 |",
        "| 初版（过拟合） | 60.90% | 显著高于验证集 | **>27%** | Dropout偏低、增强较弱 |",
    ),
    (
        "说明调优主要解决的是过拟合而非“记忆训练集”。",
        "说明调优主要解决的是过拟合，而非对训练样本的过度拟合。",
    ),
    (
        "【插入截图 图8：训练终端Epoch/Loss/Acc日志】",
        "**图8** 训练过程 Epoch/Loss/Acc 日志",
    ),
    (
        "图8中应可见Epoch 22附近验证Acc达峰、之后震荡或回落，支撑早停决策。",
        "图8显示 Epoch 22 附近验证 Acc 达到峰值，之后出现震荡或回落，支撑早停决策。",
    ),
    (
        "【插入截图 图9：训练/验证Loss与Acc曲线（据history或日志绘制）】",
        "**图9** 训练与验证 Loss、Acc 曲线",
    ),
    (
        "图9用于直观展示表8中“差距收窄”趋势。",
        "图9展示表8中训练集与验证集准确率差距随训练轮次的变化趋势。",
    ),
    (
        "【插入截图 图10：`training_config.json`与`deepfashion_best_model.pth`文件属性】",
        "**图10** 训练配置文件与模型权重文件信息",
    ),
    (
        "【插入截图 图11：测试集混淆矩阵（英文：`docs/evaluation_confusion_matrix.png`；中文：`docs/evaluation_confusion_matrix_zh.png`）】",
        "**图11** 测试集混淆矩阵",
    ),
    (
        "App展示Top-3有数据支撑。",
        "应用 Top-3 展示与测试集 Top-3 准确率 79.00% 的评估结果一致。",
    ),
    (
        "| 36 | Romper 连体短裤 | Dress 连衣裙 | 连身装互混 |",
        "| 36 | Romper 连体短裤 | Dress 连衣裙 | 连身装类别相互混淆 |",
    ),
    (
        "| 36 | Jumpsuit 连体衣 | Dress 连衣裙 | 连身装互混 |",
        "| 36 | Jumpsuit 连体衣 | Dress 连衣裙 | 连身装类别相互混淆 |",
    ),
    (
        "| 26 | Shorts 短裤 | Skirt 裙子 | 下装类互混 |",
        "| 26 | Shorts 短裤 | Skirt 裙子 | 下装类别相互混淆 |",
    ),
    (
        "【插入截图 图12：`docs/misclassified_samples/`中2～3张错例（如Romper→Jumpsuit，conf≈0.98）】",
        "**图12** 测试集高置信度错分样例",
    ),
    (
        "典型错例：Romper↔Jumpsuit（连体类互混）、Tee→Tank、Blouse→Jumpsuit等。",
        "典型错例：Romper↔Jumpsuit（连体装类别相互混淆）、Tee→Tank、Blouse→Jumpsuit 等。",
    ),
    (
        "【插入截图 图13：上述`test_onnx_model.py`终端输出】",
        "**图13** ONNX 与 PyTorch 推理一致性验证输出",
    ),
    (
        "图13证明部署图与训练权重在真实样本上数值一致，优于仅用随机张量做 smoke test。",
        "图13表明，部署模型与训练权重在真实样本上推理结果一致，优于仅使用随机张量进行的冒烟测试。",
    ),
    (
        "【插入截图 图14：根目录ONNX与Android assets/models路径对照】",
        "**图14** ONNX 模型路径对照",
    ),
    (
        "图14证明训练产出已进入APK打包路径。",
        "图14表明，训练导出的 ONNX 模型已纳入 APK 资源目录。",
    ),
    (
        "应用名称「服饰识别」v1.2，承担**模型消费与结果沉淀**，不参与训练。以下按实际页面说明各截图应呈现的内容与数据挖掘含义。",
        "应用名称「服饰识别」v1.2，负责模型推理与识别结果的本地存储，不参与训练。以下各页面对应部署验证中的关键界面与数据来源。",
    ),
    (
        "【插入截图 图15：Android Studio 中 DeepFashionClassifier 工程，展开 `app/src/main/assets/models/` 与 `java/com/deepfashion/classifier/` 核心类】",
        "**图15** Android 工程结构与模型资源路径",
    ),
    (
        "图15应突出与数据挖掘成果直接相关的路径：",
        "图15列出与模型部署直接相关的路径：",
    ),
    (
        "【插入截图 图16：拍照或相册识别后的结果页，含 Top-1 类别、置信度进度条、Top-3 Tab 与百科卡片】",
        "**图16** 识别结果页（ResultActivity）",
    ),
    (
        "3. **置信度条形图**：三个候选的横向对比条，便于用户直观比较概率差距。",
        "3. **置信度条形图**：三个候选类别的横向概率对比条。",
    ),
    (
        "图16说明：即使 Top-1 误判，用户仍可通过 Top-3 Tab 看到次优候选，降低单次错误带来的体验损失。",
        "图16中，Top-3 候选列表可补充展示次优类别及其置信度，用于在 Top-1 误判时提供备选识别结果。",
    ),
    (
        "【插入截图 图17：更多 → 数据统计，含汇总数字、7日/30日趋势柱状图、类别饼图】",
        "**图17** 数据统计页（StatisticsActivity）",
    ),
    (
        "上述统计均从本地 `history.jsonl` 读取，**不上传云端**，属于部署后的**行为型数据挖掘**：可观察用户常识别哪类服饰、何时使用频繁、平均置信度是否偏低等，为模型迭代提供侧写（与第6章测试集统计互补）。",
        "上述统计均从本地 `history.jsonl` 读取，**不上传云端**，属于基于本地使用记录的统计分析，可补充考察识别类别分布、使用频次与平均置信度等指标，与第6章基于固定测试集的离线评估形成补充。",
    ),
    (
        "【插入截图 图18：更多 → 模型信息，完整文本卡片】",
        "**图18** 模型信息页（ModelInfoActivity）",
    ),
    # docx 可能无 markdown 加粗
    (
        "图5 同一张原图经增强后的数据增强前后对比",
        "图5 数据增强前后对比",
    ),
    ("对比拼图", "数据增强前后对比"),
    ("明显更高", "显著高于验证集"),
    ("连身装互混", "连身装类别相互混淆"),
    ("下装类互混", "下装类别相互混淆"),
    (
        "3. 置信度条形图：三个候选的横向对比条，便于用户直观比较概率差距。",
        "3. 置信度条形图：三个候选类别的横向概率对比条。",
    ),
    (
        "上述统计均从本地 history.jsonl 读取，不上传云端，属于部署后的行为型数据挖掘：可观察用户常识别哪类服饰、何时使用频繁、平均置信度是否偏低等，为模型迭代提供侧写（与第6章测试集统计互补）。",
        "上述统计均从本地 history.jsonl 读取，不上传云端，属于基于本地使用记录的统计分析，可补充考察识别类别分布、使用频次与平均置信度等指标，与第6章基于固定测试集的离线评估形成补充。",
    ),
]


def apply_replacements(text: str) -> tuple[str, int]:
    count = 0
    for old, new in REPLACEMENTS:
        if old in text:
            text = text.replace(old, new)
            count += 1
    return text, count


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    t = paragraph.text.replace(old, new)
    runs = paragraph.runs
    if runs:
        runs[0].text = t
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(t)
    return True


def polish_docx() -> int:
    if Document is None or not DOCX_PATH.exists():
        return 0
    doc = Document(str(DOCX_PATH))
    total = 0
    for old, new in REPLACEMENTS:
        for paragraph in doc.paragraphs:
            if replace_in_paragraph(paragraph, old, new):
                total += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if replace_in_paragraph(paragraph, old, new):
                            total += 1
    doc.save(str(DOCX_PATH))
    return total


def main() -> None:
    if not MD_PATH.exists():
        print(f"找不到: {MD_PATH}")
        sys.exit(1)

    md_text = MD_PATH.read_text(encoding="utf-8")
    md_text, md_count = apply_replacements(md_text)
    MD_PATH.write_text(md_text, encoding="utf-8")
    print(f"Markdown 替换 {md_count} 处 → {MD_PATH}")

    n = polish_docx()
    print(f"Word 段落/单元格替换 {n} 处 → {DOCX_PATH}")


if __name__ == "__main__":
    main()
