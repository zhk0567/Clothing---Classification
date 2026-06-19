#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""最终检查《数据挖掘实训报告》md / docx / 插图一致性。"""

from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent


def main() -> int:
    sys.stdout.reconfigure(encoding="utf-8")
    issues: list[str] = []
    warnings: list[str] = []

    md = (ROOT / "数据挖掘实训报告.md").read_text(encoding="utf-8")
    md_figs = re.findall(r"\*\*(图\d+[^*]*)\*\*", md)
    md_tables = re.findall(r"\*\*(表\d+[^*]*)\*\*", md)
    md_fig_nums = {int(re.match(r"图(\d+)", f).group(1)) for f in md_figs if re.match(r"图(\d+)", f)}
    md_tbl_nums = {int(re.match(r"表(\d+)", t).group(1)) for t in md_tables if re.match(r"表(\d+)", t)}

    for i in range(1, 19):
        if i not in md_fig_nums:
            issues.append(f"MD 缺少图{i}标题")
    for i in range(1, 21):
        if i not in md_tbl_nums:
            issues.append(f"MD 缺少表{i}标题")

    for pat in [r"图\d+-\d+", r"表\d+-\d+", "（已修复）", "可答辩", "图14："]:
        if re.search(pat, md):
            issues.append(f"MD 残留模式: {pat}")

    try:
        from docx import Document
    except ImportError:
        issues.append("未安装 python-docx，跳过 docx 检查")
        doc = None
    else:
        doc = Document(str(ROOT / "数据挖掘实训报告.docx"))
        fig_titles = [p.text.strip() for p in doc.paragraphs if p.style and p.style.name == "图标题"]
        tbl_captions = [p.text.strip() for p in doc.paragraphs if p.style and p.style.name == "表题注"]
        all_para = "\n".join(p.text for p in doc.paragraphs)

        fig_nums = set()
        for t in fig_titles:
            m = re.match(r"图(\d+)", t)
            if m:
                fig_nums.add(int(m.group(1)))
        for i in range(1, 19):
            if i not in fig_nums:
                issues.append(f"DOCX 缺少图{i}标题")

        for pat in [r"图\d+-\d+", r"表\d+-\d+", "（已修复）", "可答辩", "图14：", "图6-3", "图5-2"]:
            if re.search(pat, all_para):
                issues.append(f"DOCX 残留模式: {pat}")

        if "图9展示表8" not in all_para:
            warnings.append("DOCX 图9交叉引用「图9展示表8…」未找到")

        img_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)
        if img_count < 14:
            warnings.append(f"DOCX 嵌入图片仅 {img_count} 张，可能仍缺图8/9/10/13 等待插入")

        # md vs docx title drift (first 10 chars)
        md_map = {}
        for f in md_figs:
            m = re.match(r"图(\d+)", f)
            if m:
                md_map[int(m.group(1))] = re.sub(r"^\*\*|\*\*$", "", f).strip()

        for t in fig_titles:
            m = re.match(r"图(\d+)", t)
            if not m:
                continue
            n = int(m.group(1))
            md_title = md_map.get(n, "")
            if md_title and t.split(maxsplit=1)[-1][:6] != md_title.split(maxsplit=1)[-1][:6]:
                # loose compare - only warn on obvious mismatch
                if n in {3, 8, 12, 15}:
                    warnings.append(f"图{n} 标题 md/docx 不完全一致:\n    md: {md_title}\n    docx: {t}")

    fig_dir = ROOT / "docs" / "report_figures"
    img_ext = {".png", ".jpg", ".jpeg"}
    files = [f for f in fig_dir.iterdir() if f.is_file() and f.suffix.lower() in img_ext]
    old_files = [f.name for f in files if re.match(r"图\d+-\d+", f.name)]
    if old_files:
        issues.append(f"插图目录仍有旧命名: {old_files}")

    for n in range(1, 19):
        prefix = f"图{n}_"
        if not any(f.name.startswith(prefix) for f in files):
            issues.append(f"插图目录缺少图{n} ({prefix}*)")

    # key metrics consistency
    if "61.55%" not in md or "58.73%" not in md:
        issues.append("MD 关键指标 61.55% / 58.73% 缺失")
    if "比选" not in md and "界定" not in md:
        warnings.append("MD 未找到「比选/界定50类」表述，请人工确认")

    print("=== 检查摘要 ===")
    print(f"MD: 图 {len(md_fig_nums)}/18, 表 {len(md_tbl_nums)}/20")
    if doc:
        print(f"DOCX: 图标题 {len(fig_titles)}/18, 表题注 {len(tbl_captions)}, 嵌入图 {img_count}")
    print(f"插图目录: {len(files)} 个文件")
    print()

    if warnings:
        print("=== 提示（非致命）===")
        for w in warnings:
            print(" ·", w)
        print()

    if issues:
        print("=== 需处理 ===")
        for x in issues:
            print(" ✗", x)
        return 1

    print("=== 自动检查通过 ===")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
