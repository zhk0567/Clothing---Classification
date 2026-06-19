#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""同步数据集链接与本地爬取说明到 docx（仅改文字）。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOCX = Path(__file__).resolve().parent.parent / "数据挖掘实训报告.docx"

LOCAL_OLD = (
    "项目早期建立本地目录服装/，按上衣、下装、连身装三个粗粒度文件夹存放图片，共11287张，意图构建贴近实拍场景的数据源。"
)
LOCAL_NEW = (
    "项目早期建立本地服装目录，按上衣、下装、连身装三个粗粒度文件夹存放图片，共11287张，意图构建贴近电商实拍场景的数据源。"
    "目录内图像由作者在淘宝、京东等电商平台自行爬取获得，保存的是商品展示图；"
    "数据未经逐图精细标注，仅依据文件夹名称赋予三类弱标签，未进入本课设训练管线。"
)

DEEP_OLD = "正式数据源为Category and Attribute Prediction Benchmark。"
DEEP_NEW = (
    "正式数据源为 Category and Attribute Prediction Benchmark，对应 DeepFashion 数据集的类别与属性预测任务子集。"
    "本课设使用的训练、验证与测试划分及 50 类服饰标签均来自该基准。"
    "数据集可通过 HyperAI 开源镜像获取；结构说明、任务定义与论文引用信息见香港中文大学 MMLab 官方项目页："
    "下载 https://hyper.ai/cn/datasets/16118 ；说明 https://mmlab.ie.cuhk.edu.hk/projects/DeepFashion.html 。"
)


def rep(paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    t = paragraph.text.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = t
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(t)
    return True


def insert_paragraph_after(paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    p.add_run(text)
    return p


def patch_table4(doc: Document) -> bool:
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        if header[0] == "维度" and "DeepFashion" in header[-1]:
            if header[1] == "本地服装/":
                table.rows[0].cells[1].text = "本地服装目录"
            if any(r.cells[0].text.strip() == "数据来源" for r in table.rows):
                return False
            new_row = table.add_row()
            cells = new_row.cells
            cells[0].text = "数据来源"
            cells[1].text = "淘宝、京东爬取，作者自行采集"
            cells[2].text = "HyperAI 镜像下载，见 2.2 节链接"
            # 移到表头后：python-docx add_row 在末尾，需手动调整顺序较复杂
            # 将最后一行移到 index 1
            tbl = table._tbl
            tr = new_row._tr
            tbl.remove(tr)
            tbl.insert(1, tr)
            return True
    return False


def patch_references(doc: Document) -> bool:
    refs = [
        "[8] HyperAI. DeepFashion 数据集[EB/OL]. https://hyper.ai/cn/datasets/16118.",
        "[9] MMLab CUHK. DeepFashion Project[EB/OL]. https://mmlab.ie.cuhk.edu.hk/projects/DeepFashion.html.",
    ]
    existing = "\n".join(p.text for p in doc.paragraphs)
    if refs[0] in existing:
        return False
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("[7]"):
            anchor = p
            for ref in refs:
                insert_paragraph_after(anchor, ref)
                anchor = anchor._p.getnext()
                while anchor is not None and anchor.tag.endswith("p"):
                    # find last inserted paragraph object - simpler: chain from p
                    pass
            # simpler loop
            last = p
            for ref in refs:
                np = insert_paragraph_after(last, ref)
                last = np
            return True
    return False


def main() -> None:
    doc = Document(str(DOCX))
    n = 0
    for p in doc.paragraphs:
        if rep(p, LOCAL_OLD, LOCAL_NEW):
            n += 1
        if rep(p, DEEP_OLD, DEEP_NEW):
            n += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if rep(p, LOCAL_OLD, LOCAL_NEW):
                        n += 1
                    if rep(p, DEEP_OLD, DEEP_NEW):
                        n += 1
    if patch_table4(doc):
        n += 1
    if patch_references(doc):
        n += 1
    doc.save(str(DOCX))
    print(f"已更新 {n} 处 → {DOCX}")


if __name__ == "__main__":
    main()
