#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""仅替换 Word 报告中的图/表编号文字，不改动段落与表格样式。"""

from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.table import _Cell
except ImportError:
    print("请先安装: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "数据挖掘实训报告.docx"

# 先长后短，避免 表6-1 误伤 表6-1b
REPLACEMENTS = [
    ("表6-1b", "表11"),
    ("表5-1b", "表7"),
    ("表6-3与6-4", "表13与表14"),
    ("表5-1与5-1b", "表6与表7"),
    ("表1-1", "表1"),
    ("表2-1", "表2"),
    ("表2-2", "表3"),
    ("表2-3", "表4"),
    ("表4-1", "表5"),
    ("表5-1", "表6"),
    ("表5-2", "表8"),
    ("表5-3", "表9"),
    ("表6-1", "表10"),
    ("表6-2", "表12"),
    ("表6-3", "表13"),
    ("表6-4", "表14"),
    ("表7-1", "表15"),
    ("表7-2", "表16"),
    ("表A-1", "表17"),
    ("表A-2", "表18"),
    ("表A-3", "表19"),
    ("表A-4", "表20"),
    ("图6-3", "图13"),
    ("图6-2", "图12"),
    ("图6-1", "图11"),
    ("图5-3", "图10"),
    ("图5-2", "图9"),
    ("图5-1", "图8"),
    ("图4-1", "图7"),
    ("图3-3", "图6"),
    ("图3-2", "图5"),
    ("图3-1", "图4"),
    ("图2-3", "图3"),
    ("图2-2", "图2"),
    ("图2-1", "图1"),
    ("图7-5", "图18"),
    ("图7-4", "图17"),
    ("图7-3", "图16"),
    ("图7-2", "图15"),
    ("图7-1", "图14"),
    ("（图7-5）", "图18"),
    ("（图7-4）", "图17"),
    ("（图7-3）", "图16"),
    ("（图7-2）", "图15"),
    ("（图7-1）", "图14"),
]


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    full = paragraph.text
    if old not in full:
        return False
    new_full = full.replace(old, new)
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_full)
        return True
    runs[0].text = new_full
    for run in runs[1:]:
        run.text = ""
    return True


def replace_in_cell(cell: _Cell, old: str, new: str) -> bool:
    changed = False
    for paragraph in cell.paragraphs:
        if replace_in_paragraph(paragraph, old, new):
            changed = True
    return changed


def replace_in_table(table, old: str, new: str) -> int:
    count = 0
    for row in table.rows:
        for cell in row.cells:
            if replace_in_cell(cell, old, new):
                count += 1
    return count


def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is None:
                continue
            for p in part.paragraphs:
                yield p
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p


def main() -> None:
    if not DOCX_PATH.exists():
        print(f"找不到文件: {DOCX_PATH}")
        sys.exit(1)

    doc = Document(str(DOCX_PATH))
    stats: dict[str, int] = {}

    for old, new in REPLACEMENTS:
        n = 0
        for paragraph in iter_all_paragraphs(doc):
            if replace_in_paragraph(paragraph, old, new):
                n += 1
        if n:
            stats[f"{old} -> {new}"] = n

    doc.save(str(DOCX_PATH))
    print(f"已更新: {DOCX_PATH}")
    for k, v in stats.items():
        print(f"  {k}: {v} 处")


if __name__ == "__main__":
    main()
