#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将《数据挖掘实训报告.md》转换为 Word 文档（.docx）。

用法（在项目根目录）:
    python scripts/build_training_report_docx.py

依赖:
    pip install python-docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, Cm
except ImportError:
    print("缺少 python-docx，请先执行: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "数据挖掘实训报告.md"
OUT_PATH = ROOT / "数据挖掘实训报告.docx"

SCREENSHOT_RE = re.compile(r"^【插入截图\s*(.+?)】$")
HEADING1_RE = re.compile(r"^##\s+(.+)$")
HEADING2_RE = re.compile(r"^###\s+(.+)$")
HEADING3_RE = re.compile(r"^####\s+(.+)$")
HR_RE = re.compile(r"^-{3,}$")
TABLE_SEP_RE = re.compile(r"^\|[-:\s|]+\|$")
TABLE_CAPTION_RE = re.compile(r"^\*\*(表[\dA-Za-z\-]+(?:\s+.+)?)\*\*$")


def set_run_font(run, name: str = "宋体", size_pt: float = 12, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def setup_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)


def add_paragraph(doc: Document, text: str, *, bold: bool = False, align_center: bool = False):
    p = doc.add_paragraph()
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    if level == 1:
        p.style = doc.styles["Heading 1"]
        size = 16
    elif level == 2:
        p.style = doc.styles["Heading 2"]
        size = 14
    else:
        p.style = doc.styles["Heading 3"]
        size = 13
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=True)
    return p


def parse_table_rows(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        if TABLE_SEP_RE.match(line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def set_cell_shading(cell, fill: str = "D9D9D9") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def mark_row_as_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tr_pr.append(tbl_header)


def add_table_caption(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    set_run_font(run, size_pt=10.5, bold=True)


def add_table(doc: Document, rows: list[list[str]], *, caption: str | None = None):
    if not rows:
        return
    if caption:
        add_table_caption(doc, caption)
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        if r_idx == 0:
            mark_row_as_table_header(table.rows[0])
        for c_idx in range(cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell_text = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
            cell_text = re.sub(r"`(.+?)`", r"\1", cell_text)
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(cell_text)
            is_header = r_idx == 0
            set_run_font(run, size_pt=10.5, bold=is_header)
            if is_header:
                set_cell_shading(cell)


def add_codeblock(doc: Document, code_lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    text = "\n".join(code_lines)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size_pt=10)


def convert_md_to_docx(md_path: Path, out_path: Path) -> None:
    if not md_path.exists():
        raise FileNotFoundError(f"找不到 Markdown 文件: {md_path}")

    content = md_path.read_text(encoding="utf-8")
    lines = content.splitlines()

    doc = Document()
    setup_document_styles(doc)

    i = 0
    pending_table_caption: str | None = None
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if not line.strip():
            i += 1
            continue

        if line.startswith("# ") and not line.startswith("## "):
            add_heading(doc, line[2:].strip(), level=1)
            i += 1
            continue

        m1 = HEADING1_RE.match(line)
        if m1:
            add_heading(doc, m1.group(1).strip(), level=1)
            i += 1
            continue

        m2 = HEADING2_RE.match(line)
        if m2:
            add_heading(doc, m2.group(1).strip(), level=2)
            i += 1
            continue

        m3 = HEADING3_RE.match(line)
        if m3:
            add_heading(doc, m3.group(1).strip(), level=3)
            i += 1
            continue

        if HR_RE.match(line.strip()):
            i += 1
            continue

        if line.strip().startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            add_codeblock(doc, code_lines)
            continue

        cap_match = TABLE_CAPTION_RE.match(line.strip())
        if cap_match:
            pending_table_caption = cap_match.group(1).strip()
            i += 1
            continue

        if line.strip().startswith("|"):
            rows, i = parse_table_rows(lines, i)
            add_table(doc, rows, caption=pending_table_caption)
            pending_table_caption = None
            continue

        ms = SCREENSHOT_RE.match(line.strip())
        if ms:
            caption = ms.group(1).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"【插入截图 {caption}】")
            set_run_font(run, bold=True, size_pt=11)
            i += 1
            continue

        if line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line.strip()[2:])
            set_run_font(run)
            i += 1
            continue

        text = line.strip()
        if text.startswith("**") and text.endswith("**"):
            add_paragraph(doc, text.strip("*").strip(), bold=True)
        elif line.startswith("**") and line.endswith("**"):
            add_paragraph(doc, line.strip("*"), bold=True, align_center="封面" in lines[max(0, i - 5): i + 1][0] if False else False)
        else:
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            clean = re.sub(r"`(.+?)`", r"\1", clean)
            center = any(k in clean for k in ("学 院", "专 业", "学 号", "姓 名", "题 目", "指导教师", "日 期"))
            add_paragraph(doc, clean, align_center=center)

        i += 1

    doc.save(out_path)
    print(f"已生成: {out_path}")


def main():
    print(f"读取: {MD_PATH}")
    convert_md_to_docx(MD_PATH, OUT_PATH)


if __name__ == "__main__":
    main()
